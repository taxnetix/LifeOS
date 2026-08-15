"""DOCX extraction — wills, trust deeds, agreements.

Module is `docx_` because `python-docx` imports as `docx`; a sibling named
`docx.py` would shadow it.

Paragraph numbering is 1-based over NON-EMPTY paragraphs only, so the locator
stays stable when someone opens the document in Word and it re-saves with
different whitespace. An unstable locator would change record ids and break
idempotency.
"""

from __future__ import annotations

from pathlib import Path

from . import Block, Extraction, Table

_CONF = 0.99  # a docx has a real text layer; nothing is being inferred


def extract_docx(path: Path, doc_hash: str) -> Extraction:
    import docx

    out = Extraction(doc_hash=doc_hash, method="python-docx", pages=1)
    document = docx.Document(str(path))

    try:
        props = document.core_properties
        out.meta = {
            k: str(v)
            for k, v in {
                "title": props.title,
                "author": props.author,
                "created": props.created,
                "modified": props.modified,
            }.items()
            if v
        }
    except Exception:  # noqa: BLE001 — metadata is a nicety, never a blocker
        pass

    n = 0
    for para in document.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        n += 1
        style = (para.style.name or "").lower() if para.style else ""
        out.blocks.append(
            Block(
                locator=f"para={n};style={style}" if "heading" in style else f"para={n}",
                text=text,
                confidence=_CONF,
            )
        )

    for t, table in enumerate(document.tables, 1):
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        rows = [r for r in rows if any(r)]
        if rows:
            out.tables.append(Table(locator=f"table={t}", rows=rows, confidence=_CONF))

    if not out.blocks and not out.tables:
        out.errors.append("document contains no text or tables")
    return out

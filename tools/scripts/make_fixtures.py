"""Generate the FAKE documents in vault.example/inbox/.

Every person, number and institution here is invented. This is the corpus every
test and demo runs against, so nothing real ever has to be.
"""
from __future__ import annotations
import csv, subprocess, sys
from pathlib import Path

ROOT = Path(__file__).resolve().parents[2]
INBOX = ROOT / "vault.example" / "inbox"
INBOX.mkdir(parents=True, exist_ok=True)

# ── 1. Bank statement PDF (text layer) ───────────────────────────────────────
ROWS = [
    ("01 Jul 2026", "OPENING BALANCE", "", "", "42 118.55"),
    ("02 Jul 2026", "SALARY NORTHWIND SOFTWARE", "", "68 420.00", "110 538.55"),
    ("03 Jul 2026", "DEBIT ORDER OUTSURANCE 4471", "1 842.00", "", "108 696.55"),
    ("03 Jul 2026", "DEBIT ORDER DISCOVERY HEALTH", "7 314.00", "", "101 382.55"),
    ("05 Jul 2026", "BOND REPAYMENT HOMELOAN 9921", "18 750.00", "", "82 632.55"),
    ("07 Jul 2026", "WOOLWORTHS SANDTON", "1 284.31", "", "81 348.24"),
    ("09 Jul 2026", "NETFLIX.COM", "199.00", "", "81 149.24"),
    ("09 Jul 2026", "SPOTIFY AB", "119.99", "", "81 029.25"),
    ("12 Jul 2026", "ESKOM PREPAID", "1 500.00", "", "79 529.25"),
    ("15 Jul 2026", "RENT RECEIVED JHB PROPS CC", "", "9 500.00", "89 029.25"),
    ("18 Jul 2026", "SARS EFILING IRP6", "12 400.00", "", "76 629.25"),
    ("22 Jul 2026", "VIRGIN ACTIVE", "899.00", "", "75 730.25"),
    ("25 Jul 2026", "ALLAN GRAY RA CONTRIB", "6 500.00", "", "69 230.25"),
    ("28 Jul 2026", "SERVICE FEE", "121.00", "", "69 109.25"),
    ("31 Jul 2026", "CLOSING BALANCE", "", "", "69 109.25"),
]
rows_html = "".join(
    f"<tr><td>{d}</td><td>{desc}</td><td class=n>{out}</td><td class=n>{inn}</td><td class=n>{bal}</td></tr>"
    for d, desc, out, inn, bal in ROWS
)
stmt = f"""<style>
@page {{ size: A4; margin: 18mm; }}
body {{ font-family: Helvetica, Arial, sans-serif; font-size: 9.5pt; }}
h1 {{ font-size: 15pt; margin: 0 0 2mm; }}
.hdr {{ border-bottom: 2px solid #333; padding-bottom: 3mm; margin-bottom: 5mm; }}
.meta td {{ padding: 0.6mm 6mm 0.6mm 0; font-size: 9pt; }}
table.tx {{ width: 100%; border-collapse: collapse; margin-top: 4mm; }}
table.tx th {{ text-align: left; border-bottom: 1px solid #333; padding: 1.5mm 1mm; font-size: 8.5pt; }}
table.tx td {{ padding: 1.2mm 1mm; border-bottom: 0.3px solid #ddd; }}
td.n, th.n {{ text-align: right; }}
.foot {{ margin-top: 6mm; font-size: 7.5pt; color: #666; }}
</style>
<div class=hdr>
  <h1>NORTHBANK</h1>
  <div>Cheque Account Statement</div>
</div>
<table class=meta>
  <tr><td><b>Account holder</b></td><td>A SAMPLE</td>
      <td><b>Statement period</b></td><td>01 Jul 2026 &ndash; 31 Jul 2026</td></tr>
  <tr><td><b>Account number</b></td><td>1049 5560 118</td>
      <td><b>Branch code</b></td><td>250655</td></tr>
  <tr><td><b>Account type</b></td><td>Gold Cheque</td>
      <td><b>Statement no.</b></td><td>087</td></tr>
</table>
<table class=tx>
  <tr><th>Date</th><th>Description</th><th class=n>Money out</th>
      <th class=n>Money in</th><th class=n>Balance</th></tr>
  {rows_html}
</table>
<div class=foot>FICTIONAL DOCUMENT — generated for LifeOS tests. No real person or account.</div>
"""

# ── 2. Policy schedule PDF ───────────────────────────────────────────────────
policy = """<style>
@page { size: A4; margin: 20mm; }
body { font-family: Georgia, serif; font-size: 10pt; }
h1 { font-size: 14pt; } h2 { font-size: 11pt; margin-top: 6mm; }
table { border-collapse: collapse; width: 100%; margin-top: 3mm; }
td { padding: 1.5mm 3mm 1.5mm 0; vertical-align: top; }
td.k { width: 45%; color: #444; }
.foot { margin-top: 8mm; font-size: 7.5pt; color: #666; }
</style>
<h1>SOUTHERN MUTUAL ASSURANCE</h1>
<div>Life Policy Schedule</div>
<h2>Policy details</h2>
<table>
  <tr><td class=k>Policy number</td><td>SM-4471902</td></tr>
  <tr><td class=k>Policy type</td><td>Whole life assurance</td></tr>
  <tr><td class=k>Life assured</td><td>A Sample</td></tr>
  <tr><td class=k>Policy owner</td><td>A Sample</td></tr>
  <tr><td class=k>Inception date</td><td>01 March 2019</td></tr>
  <tr><td class=k>Anniversary date</td><td>01 March</td></tr>
  <tr><td class=k>Sum assured</td><td>R 4 500 000.00</td></tr>
  <tr><td class=k>Monthly premium</td><td>R 1 842.00</td></tr>
  <tr><td class=k>Premium escalation</td><td>5% per annum on anniversary</td></tr>
</table>
<h2>Beneficiary nomination</h2>
<table>
  <tr><td class=k>M Sample (spouse)</td><td>60%</td></tr>
  <tr><td class=k>The Sample Family Trust</td><td>40%</td></tr>
</table>
<h2>Cession</h2>
<table><tr><td class=k>Ceded to</td><td>Northbank Home Loans as security for bond 9921</td></tr></table>
<h2>Exclusions</h2>
<table>
  <tr><td class=k>Suicide</td><td>First 24 months from inception</td></tr>
  <tr><td class=k>Aviation</td><td>Non-commercial piloting excluded</td></tr>
</table>
<div class=foot>FICTIONAL DOCUMENT — generated for LifeOS tests.</div>
"""

def render(html: str, out: Path) -> None:
    from weasyprint import HTML
    HTML(string=html).write_pdf(str(out))
    print("  ", out.name)

print("PDFs:")
render(stmt, INBOX / "northbank-cheque-jul-2026.pdf")
render(policy, INBOX / "southern-mutual-policy-schedule.pdf")

# ── 3. Scanned document (image, forces the OCR path) ─────────────────────────
print("scanned image:")
from PIL import Image, ImageDraw, ImageFont
img = Image.new("RGB", (1654, 700), "white")
d = ImageDraw.Draw(img)
try:
    big = ImageFont.truetype("/System/Library/Fonts/Supplemental/Arial Bold.ttf", 46)
    reg = ImageFont.truetype("/System/Library/Fonts/Supplemental/Arial.ttf", 34)
except OSError:
    big = reg = ImageFont.load_default()
d.text((70, 55), "DEED OF GRAVE", font=big, fill="black")
for i, line in enumerate([
    "Cemetery: Fernwood Memorial Park",
    "Section: C   Plot number: 214",
    "Registered holder: A Sample",
    "Date of purchase: 14 November 2021",
    "Receipt number: FMP-2021-8842",
]):
    d.text((70, 155 + i * 58), line, font=reg, fill="black")
d.text((70, 630), "FICTIONAL DOCUMENT - LifeOS test fixture", font=reg, fill="grey")
img.save(INBOX / "deed-of-grave-scan.png", dpi=(300, 300))
print("   deed-of-grave-scan.png")

# ── 4. CSV — a card export with a semicolon delimiter, as SA banks emit ──────
print("csv:")
with (INBOX / "capital-card-jul-2026.csv").open("w", newline="", encoding="utf-8") as fh:
    w = csv.writer(fh, delimiter=";")
    w.writerow(["Date", "Description", "Amount", "Balance"])
    for r in [
        ["2026-07-02", "TAKEALOT.COM", "-1249.00", "-8420.11"],
        ["2026-07-06", "UBER TRIP", "-142.50", "-8562.61"],
        ["2026-07-11", "APPLE.COM/BILL", "-499.00", "-9061.61"],
        ["2026-07-19", "PAYMENT RECEIVED", "9061.61", "0.00"],
    ]:
        w.writerow(r)
print("   capital-card-jul-2026.csv")

# ── 5. XLSX — an investment platform statement ───────────────────────────────
print("xlsx:")
from openpyxl import Workbook
wb = Workbook(); ws = wb.active; ws.title = "Holdings"
ws.append(["Platform", "Account", "Fund", "Units", "Price", "Value", "TER %"])
for row in [
    ["Meridian Invest", "MI-88421", "Global Equity Feeder", 1420.55, 42.18, 59918.79, 0.85],
    ["Meridian Invest", "MI-88421", "SA Bond Fund", 8800.00, 12.04, 105952.00, 0.55],
    ["Meridian Invest", "MI-88422", "Tax Free Savings", 2100.00, 31.77, 66717.00, 0.40],
]:
    ws.append(row)
ws2 = wb.create_sheet("Summary")
ws2["A1"] = "Total value"; ws2["B1"] = 232587.79
ws2["A2"] = "As at";       ws2["B2"] = "2026-07-31"
wb.save(INBOX / "meridian-holdings-jul-2026.xlsx")
print("   meridian-holdings-jul-2026.xlsx")

# ── 6. DOCX — a will ─────────────────────────────────────────────────────────
print("docx:")
import docx as dx
doc = dx.Document()
doc.add_heading("LAST WILL AND TESTAMENT", 0)
doc.add_paragraph("of A SAMPLE (Identity number 8801015800086)")
doc.add_heading("1. Revocation", level=1)
doc.add_paragraph("I revoke all previous wills and codicils made by me.")
doc.add_heading("2. Appointment of executor", level=1)
doc.add_paragraph("I nominate Meyer & Partners Attorneys of Pretoria as executor, "
                  "and direct that they be exempted from furnishing security.")
doc.add_heading("3. Guardianship", level=1)
doc.add_paragraph("Should my spouse predecease me, I nominate my sister, L Sample, "
                  "as guardian of my minor children.")
doc.add_heading("4. Bequests", level=1)
doc.add_paragraph("I bequeath my entire estate to my spouse, M Sample. Should she "
                  "predecease me, the residue shall devolve upon the Sample Family Trust "
                  "(Master's reference IT 4471/2019).")
doc.add_heading("5. Signature", level=1)
doc.add_paragraph("Signed at Pretoria on 12 March 2023 in the presence of the "
                  "undersigned witnesses.")
t = doc.add_table(rows=3, cols=2)
for i, (a, b) in enumerate([("Testator", "A Sample"),
                            ("Witness 1", "J Botha"),
                            ("Witness 2", "P Naidoo")]):
    t.rows[i].cells[0].text = a; t.rows[i].cells[1].text = b
doc.add_paragraph("FICTIONAL DOCUMENT — LifeOS test fixture.")
doc.save(INBOX / "will-a-sample-2023.docx")
print("   will-a-sample-2023.docx")

print(f"\n{len(list(INBOX.iterdir()))} fixtures in vault.example/inbox/")
